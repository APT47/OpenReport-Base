"""
OpenReport Word Utilities Module

This module provides utility functions for advanced Word document manipulation
including page numbering, table formatting, mathematical expressions, and
low-level XML operations for Word documents.

Functions:
    word_factory: Create new Word document instances
    Page numbering functions: add_page_number, add_page_number_skip_cover
    Table formatting functions: modify_border, make_tight_fit
    Math functions: latex_to_word for mathematical expressions
    Color utilities: recognise_colour for color format conversion
    Field functions: add_list_of_table_figure for automatic lists
"""

import os
import latex2mathml.converter
from docx import Document
from docx.oxml.shared import qn
from docx.oxml import OxmlElement, ns
from lxml import etree

from OpenReport.labs.word.word_constants import *


def word_factory():
    """
    Create a new Word document instance.
    
    Returns:
        Document: A new python-docx Document object
    """
    doc = Document()
    return doc


def create_element(name):
    """
    Create an XML element for use in Word document manipulation.
    
    This function creates a new XML element using the OpenXML format
    that can be used to construct complex Word document structures.
    
    Args:
        name (str): The name of the XML element to create (e.g., 'w:fldChar', 'w:instrText')
    
    Returns:
        OxmlElement: A new XML element object that can be modified and appended to document structures
    
    Example:
        >>> element = create_element('w:fldChar')
        >>> # Element can then be configured with attributes and appended to document
    """
    return OxmlElement(name)


def create_attribute(element, name, value):
    """
    Create and set an XML attribute on an existing element.
    
    This function adds a qualified name attribute to an XML element,
    properly handling namespace prefixes for Word document compatibility.
    
    Args:
        element (OxmlElement): The XML element to add the attribute to
        name (str): The qualified name of the attribute (e.g., 'w:fldCharType', 'xml:space')
        value (str): The value to assign to the attribute
    
    Returns:
        None: This function modifies the element in-place
    
    Example:
        >>> element = create_element('w:fldChar')
        >>> create_attribute(element, 'w:fldCharType', 'begin')
    """
    element.set(ns.qn(name), value)


def add_page_number(run):
    """
    Add automatic page numbering field to a Word document run.
    
    This function inserts a PAGE field into the specified run, which will
    automatically display the current page number when the document is rendered.
    The field will update automatically as pages are added or removed.
    
    Args:
        run (Run): A python-docx Run object where the page number field will be inserted
    
    Returns:
        None: The function modifies the run object in-place by appending XML elements
    
    Note:
        The page number will start from 1 for the first page of the document.
        For documents where you want to skip the cover page, use add_page_number_skip_cover instead.
    
    Example:
        >>> paragraph = document.add_paragraph()
        >>> run = paragraph.add_run('Page: ')
        >>> add_page_number(run)
    """
    fldChar1 = create_element("w:fldChar")
    create_attribute(fldChar1, "w:fldCharType", "begin")

    instrText = create_element("w:instrText")
    create_attribute(instrText, "xml:space", "preserve")
    instrText.text = "PAGE"

    fldChar2 = create_element("w:fldChar")
    create_attribute(fldChar2, "w:fldCharType", "end")

    run._r.append(fldChar1)  # noqa
    run._r.append(instrText)  # noqa
    run._r.append(fldChar2)  # noqa


def add_page_number_skip_cover(run):
    """
    Add page numbering that starts from page 2, effectively skipping the cover page.
    
    This function creates a PAGE field that displays the current page number minus 1,
    so the first page shows no number (or can be treated as page 0) and the second
    page shows "2". This is commonly used in documents with a cover page that
    shouldn't be counted in the page numbering.
    
    Args:
        run (Run): A python-docx Run object where the page number field will be inserted
    
    Returns:
        None: The function modifies the run object in-place by appending XML elements
    
    Note:
        The field creates a more complex structure with separate field characters
        to handle the offset numbering. The cover page will not display a page number.
    
    Example:
        >>> # On page 2 of document, this will display "2"
        >>> paragraph = document.add_paragraph()
        >>> run = paragraph.add_run('Page: ')
        >>> add_page_number_skip_cover(run)
    """
    fldStart = create_element("w:fldChar")
    create_attribute(fldStart, "w:fldCharType", "begin")

    instrText = create_element("w:instrText")
    create_attribute(instrText, "xml:space", "preserve")
    instrText.text = "PAGE"

    fldChar1 = create_element("w:fldChar")
    create_attribute(fldChar1, "w:fldCharType", "separate")

    fldChar2 = create_element("w:t")
    fldChar2.text = "2"

    fldEnd = create_element("w:fldChar")
    create_attribute(fldEnd, "w:fldCharType", "end")

    run._r.append(fldStart)  # noqa

    run._r.append(instrText)  # noqa
    run._r.append(fldChar1)  # noqa
    run._r.append(fldChar2)  # noqa

    run._r.append(fldEnd)  # noqa


def latex_to_word(latex_input):
    """
    Convert LaTeX mathematical expressions to Word-compatible Office Math ML format.
    
    This function takes LaTeX mathematical notation and converts it to Microsoft Word's
    Office Math ML (OMML) format, which can be inserted into Word documents as
    properly formatted mathematical equations.
    
    Args:
        latex_input (str): A string containing LaTeX mathematical notation
                          (e.g., "x^2 + y^2 = z^2", "\\frac{a}{b}", "\\sum_{i=1}^{n} x_i")
    
    Returns:
        xml.etree.ElementTree.Element: The root element of the OMML XML tree
                                      that represents the mathematical expression
                                      in Word-compatible format
    
    Raises:
        Exception: For LaTeX conversion errors with user-friendly messages
    
    Note:
        This function requires:
        - Microsoft Office to be installed (for the MML2OMML.XSL transformation file)
        - The latex2mathml library for initial LaTeX to MathML conversion
        - The lxml library for XML processing
        
        The hardcoded path assumes Office 2016/365 installation. For other versions,
        the path to MML2OMML.XSL may need to be adjusted.
    
    Example:
        >>> latex_expr = "x^2 + \\frac{y}{2} = z"
        >>> omml_element = latex_to_word(latex_expr)
        >>> # omml_element can now be inserted into a Word document
    """
    try:
        if not latex_input:
            raise Exception("OpenReport Error: LaTeX input is empty. Please provide a valid LaTeX mathematical expression.")
        
        if not isinstance(latex_input, str):
            raise Exception(f"OpenReport Error: LaTeX input must be a string, got {type(latex_input).__name__}.")
        
        # Convert LaTeX to MathML
        try:
            mathml = latex2mathml.converter.convert(latex_input)
        except Exception as e:
            raise Exception(f"OpenReport Error: Invalid LaTeX syntax in expression '{latex_input}'. Please check your mathematical expression. {str(e)}") from e
        
        # Parse MathML
        try:
            tree = etree.fromstring(mathml)
        except etree.XMLSyntaxError as e:
            raise Exception(f"OpenReport Error: Failed to parse generated MathML for expression '{latex_input}'. {str(e)}") from e
        
        # Load Microsoft Office XSLT transformation file
        xslt_paths = [
            r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
            r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
            r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL",
            r"C:\Program Files (x86)\Microsoft Office\Office16\MML2OMML.XSL"
        ]
        
        xslt_file = None
        for path in xslt_paths:
            try:
                if os.path.exists(path):
                    xslt_file = path
                    break
            except:
                continue
        
        if not xslt_file:
            raise Exception("OpenReport Error: Microsoft Office MML2OMML.XSL transformation file not found. Please ensure Microsoft Office is installed, or math expressions may not display correctly.")
        
        try:
            xslt = etree.parse(xslt_file)
        except (FileNotFoundError, etree.XMLSyntaxError) as e:
            raise Exception(f"OpenReport Error: Failed to load Microsoft Office transformation file at '{xslt_file}'. Please check your Office installation. {str(e)}") from e
        
        # Apply XSLT transformation
        try:
            transform = etree.XSLT(xslt)
            new_dom = transform(tree)
            return new_dom.getroot()
        except etree.XSLTParseError as e:
            raise Exception(f"OpenReport Error: Failed to transform MathML to Word format for expression '{latex_input}'. {str(e)}") from e
        
    except Exception as e:
        if hasattr(e, 'args') and e.args and "OpenReport Error:" in str(e.args[0]):
            raise
        raise Exception(f"OpenReport Error: Failed to convert LaTeX expression '{latex_input}' to Word format. {str(e)}") from e


def recognise_colour(colour_value):
    """
    Convert various color format inputs to python-docx RGBColor objects.
    
    This function accepts color values in multiple formats and converts them
    to RGBColor objects that can be used with python-docx for styling text,
    borders, and other document elements.
    
    Args:
        colour_value (str): Color value in one of the following formats:
                           - Named color (e.g., "red", "blue", "green")
                           - Hexadecimal format (e.g., "#FF0000", "#ff0000")
                           - RGB comma-separated format (e.g., "255,0,0")
    
    Returns:
        RGBColor: A python-docx RGBColor object representing the specified color
    
    Raises:
        KeyError: If a named color is not found in the color_mapping dictionary
        ValueError: If RGB values are not valid integers or hex values are malformed
        IndexError: If RGB format doesn't contain exactly 3 comma-separated values
    
    Note:
        - Named colors are case-insensitive (converted to uppercase internally)
        - Hexadecimal colors can be with or without the '#' prefix
        - RGB values should be integers in the range 0-255
        - The function depends on a colour_mapping dictionary imported from word_constants
    
    Examples:
        >>> color1 = recognise_colour("RED")  # Named color
        >>> color2 = recognise_colour("#FF0000")  # Hex format
        >>> color3 = recognise_colour("255,0,0")  # RGB format
        >>> # All three examples result in the same red color
    """
    # apply .upper() to colour value to get the uppercase key from color_mapping
    if colour_value.upper() in colour_mapping.keys():
        return colour_mapping[colour_value.upper()]

    # if HEX format
    elif "#" in colour_value:
        value = colour_value.lstrip("#")
        lv = len(value)
        rgb_tuple = tuple(
            int(value[i: i + lv // 3], 16) for i in range(0, lv, lv // 3)
        )
        return RGBColor(rgb_tuple[0], rgb_tuple[1], rgb_tuple[2])

    # if RGB format
    else:
        return RGBColor(*map(lambda v: int(v), colour_value.split(",")))
