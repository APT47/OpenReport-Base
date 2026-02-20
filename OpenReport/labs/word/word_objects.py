"""
OpenReport Word Objects Module

This module contains the core classes for creating and formatting Microsoft Word document
elements. It provides abstract base classes and concrete implementations for various
document components like text, headings, math expressions, and more.
"""

from __future__ import annotations
from abc import ABC, abstractmethod

import re
import random
import string

from OpenReport.labs.word.word_attributes import *
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from OpenReport.enumerations.attribute_nodes import ATTRIBUTE_NODES
from OpenReport.enumerations.text_modes import TEXT_MODES


class WordObject(ABC):
    """
    Abstract base class for all Word document objects.
    """

    def __init__(self):
        self.document = None
        self.this_object = None
        self.attributes = None
        self.cell_attributes = None
        self.bullet_list_attributes = None
        self.caption_attributes = None

    def _prepare_document(self, word_document):
        self.document = word_document.document

    @abstractmethod
    def add_to_document(self, word_document) -> None:
        pass

    @staticmethod
    def get_attribute_class(attribute):
        return attribute_to_class_map[attribute]

    @staticmethod
    def get_run_attribute_class(run_attribute):
        return run_attribute_to_class_map[run_attribute]

    def get_document_sections_page_width(self):
        return (
            self.document.sections[0].page_width
            - self.document.sections[0].left_margin
            - self.document.sections[0].right_margin
        )

    def _apply_attributes(self):
        for kwarg, value in self.attributes.items():
            attribute_class = self.get_attribute_class(kwarg)
            attribute_class.apply(self, value)

    def _apply_cell_attributes(self):
        if self.cell_attributes != {}:
            for kwarg, value in self.cell_attributes.items():
                attribute_class = self.get_attribute_class(kwarg)
                attribute_class.apply(self.this_object, value)

    def _add_paragraph_object(self):
        # if bullet_list_attributes is specified then add the paragraph to the bullet list
        if self.bullet_list_attributes != {}:
            self.this_object = self.document.add_paragraph(
                style=self.bullet_list_attributes[ATTRIBUTE_NODES.BULLET_LIST_STYLE]
            )
        # else add the paragraph to the plain document
        else:
            self.this_object = self.document.add_paragraph()


class Paragraph(WordObject):
    """
    Base class for paragraph-based Word document objects.
    """

    def __init__(self) -> None:
        super().__init__()
        self.text = None
        self.text_attr_pairs = []
        self.object_runs = []
        self.object_fonts = []
        self.object_alignment = None
        self.paragraph_format = None

    def _prepare_and_apply_attributes(self):
        self._prepare_run_attribute_pairs()
        self._apply_runs_attributes()
        self._prepare_paragraph_attributes()
        self._apply_attributes()
        self._apply_cell_attributes()

    def _prepare_paragraph_attributes(self):
        self.object_runs = self.this_object.runs
        self.object_fonts = [run.font for run in self.object_runs]
        self.object_alignment = self.this_object.alignment
        self.paragraph_format = self.this_object.paragraph_format

    def _prepare_text(self):
        try:
            if not hasattr(self, 'attributes') or not self.attributes:
                raise Exception("OpenReport Error: Text object missing attributes. Please provide text content.")

            if ATTRIBUTE_NODES.BODY in self.attributes:
                text_source = self.attributes[ATTRIBUTE_NODES.BODY]
                if text_source is None:
                    raise Exception("OpenReport Error: Text body is None. Please provide text content.")
                self.text = str(text_source)
            else:
                raise Exception("OpenReport Error: Text object missing 'body' attribute. Please provide text content using the 'body' key.")

            if not isinstance(self.text, str):
                raise Exception(f"OpenReport Error: Text content must be a string, got {type(self.text).__name__}.")

            if not self.text:
                raise Exception("OpenReport Error: Text content is empty. Please provide non-empty text.")

            if len(self.text) >= 2:
                if self.text[0] in ("'", '"'):
                    self.text = self.text[1:]
                if self.text and self.text[-1] in ("'", '"'):
                    self.text = self.text[:-1]

            self.text = r"\normal{" + self.text + "}"

        except Exception as e:
            if hasattr(e, 'args') and e.args and "OpenReport Error:" in str(e.args[0]):
                raise
            raise Exception(f"OpenReport Error: Failed to prepare text content. {str(e)}") from e

    def _apply_runs_attributes(self):
        try:
            if not hasattr(self, 'text_attr_pairs') or not self.text_attr_pairs:
                return

            for i, pair in enumerate(self.text_attr_pairs):
                try:
                    run_text = self._get_text_from_pair(pair)
                    run_attributes = pair[run_text]

                    if TEXT_MODES.MATH_MODE in run_attributes:
                        try:
                            word_math = latex_to_word(run_text[1:])
                            self.this_object._element.append(word_math)  # noqa
                        except Exception as e:
                            raise Exception(f"OpenReport Error: Failed to convert LaTeX math expression '{run_text[1:]}' to Word format. Please check your math syntax. {str(e)}") from e

                    else:
                        try:
                            run = self.this_object.add_run(text=run_text)
                            self.this_object.runs.append(run)

                            for run_attribute in run_attributes:
                                try:
                                    attribute_class = self.get_run_attribute_class(run_attribute)
                                    attribute_class.apply(run)
                                except KeyError:
                                    raise Exception(f"OpenReport Error: Unknown text formatting attribute '{run_attribute}'. Please check your text formatting commands.")
                                except Exception as e:
                                    raise Exception(f"OpenReport Error: Failed to apply text formatting '{run_attribute}'. {str(e)}") from e

                        except Exception as e:
                            if hasattr(e, 'args') and e.args and "OpenReport Error:" in str(e.args[0]):
                                raise
                            raise Exception(f"OpenReport Error: Failed to create text run for '{run_text}'. {str(e)}") from e

                except Exception as e:
                    if hasattr(e, 'args') and e.args and "OpenReport Error:" in str(e.args[0]):
                        raise
                    raise Exception(f"OpenReport Error: Failed to process text formatting for text segment {i}. {str(e)}") from e

        except Exception as e:
            if hasattr(e, 'args') and e.args and "OpenReport Error:" in str(e.args[0]):
                raise
            raise Exception(f"OpenReport Error: Failed to apply text formatting attributes. {str(e)}") from e

    def _prepare_run_attribute_pairs(self):
        current_run_string = ""
        run_attributes = []
        char_index = -1
        command = ""

        for char in self.text:
            char_index = char_index + 1

            if TEXT_MODES.MATH_MODE in run_attributes:
                if self.text[char_index::].startswith(r"$"):
                    self.text_attr_pairs.append(
                        {current_run_string: [TEXT_MODES.MATH_MODE]}
                    )
                    run_attributes = run_attributes[:-1]
                    current_run_string = ""
                else:
                    current_run_string = current_run_string + char

            else:
                if char == "$" and command != TEXT_MODES.MATH_MODE:
                    command = TEXT_MODES.MATH_MODE
                    current_run_string = current_run_string.split("\\" + command)[0]
                    self.text_attr_pairs.append(
                        {current_run_string: list(run_attributes)}
                    )
                    run_attributes.append(command)
                    current_run_string = ""

                elif char == "$" and command == TEXT_MODES.MATH_MODE:
                    command = ""

                elif char == "{":
                    command = current_run_string.split("\\")[-1]
                    current_run_string = current_run_string.split("\\" + command)[0]
                    self.text_attr_pairs.append(
                        {current_run_string: list(run_attributes)}
                    )
                    run_attributes.append(command)
                    current_run_string = ""

                elif char == "}":
                    self.text_attr_pairs.append(
                        {current_run_string: list(run_attributes)}
                    )
                    current_run_string = ""
                    run_attributes = run_attributes[:-1]

                else:
                    current_run_string = current_run_string + char

        self.text_attr_pairs = self.text_attr_pairs[1:]

    @staticmethod
    def _get_text_from_pair(text_attr_pair) -> list:
        return list(text_attr_pair.keys())[0]

    def add_to_document(self, word_document) -> None:
        pass


class Text(Paragraph):
    """
    Text element for Word documents.
    """

    def __init__(
        self, attributes: dict, cell_attributes=None, bullet_list_attributes=None
    ) -> None:
        super().__init__()
        self.cell_attributes = cell_attributes if cell_attributes else {}
        self.bullet_list_attributes = (
            bullet_list_attributes if bullet_list_attributes else {}
        )
        self.attributes = attributes

    def add_to_document(self, word_document) -> None:
        self._prepare_document(word_document=word_document)
        self._prepare_text()
        self._add_paragraph_object()
        self._prepare_and_apply_attributes()


class MathExpression(Paragraph):
    """
    Mathematical expression element for Word documents.
    """

    def __init__(
        self, attributes: dict, bullet_list_attributes=None, cell_attributes=None
    ) -> None:
        super().__init__()
        self.attributes = attributes
        self.cell_attributes = cell_attributes if cell_attributes else {}
        self.bullet_list_attributes = (
            bullet_list_attributes if bullet_list_attributes else {}
        )

    def add_to_document(self, word_document):
        self._prepare_document(word_document=word_document)
        self._prepare_math_expression()
        self._add_paragraph_object()
        self._add_math_expression()
        self._apply_style_attributes()

    def _apply_style_attributes(self):
        unique_id = "".join(random.choices(string.ascii_letters, k=25))
        unique_meth_expression_style = self.document.styles.add_style(
            unique_id, WD_STYLE_TYPE.PARAGRAPH
        )
        unique_meth_expression_style.base_style = self.document.styles[
            self.this_object.style.name
        ]
        self.this_object.style = unique_meth_expression_style

        for kwarg, value in self.attributes.items():
            attribute_class = math_expression_attribute_to_class_map[kwarg]
            attribute_class.apply(self.this_object.style, value)

    def _prepare_math_expression(self):
        self.text = self.attributes[ATTRIBUTE_NODES.BODY]
        self.text = r"" + self.text + ""

    def _add_math_expression(self):
        word_math = latex_to_word(self.text)
        self.this_object._element.append(word_math)  # noqa


class Heading(Paragraph):
    """
    Heading element for Word documents.
    """

    def __init__(self, attributes, cell_attributes=None) -> None:
        super().__init__()
        self.heading_level = DEFAULT_HEADING_LEVEL
        self.attributes = attributes
        self.cell_attributes = cell_attributes if cell_attributes else {}

    def _prepare_heading_level(self) -> None:
        if ATTRIBUTE_NODES.LEVEL in self.attributes:
            self.heading_level = self.attributes[ATTRIBUTE_NODES.LEVEL]

    def _add_heading_object(self) -> None:
        self.this_object = self.document.add_heading(level=self.heading_level)

    def add_to_document(self, word_document) -> None:
        self._prepare_document(word_document=word_document)
        self._prepare_text()
        self._prepare_heading_level()
        self._add_heading_object()
        self._prepare_and_apply_attributes()


class DocumentParams(WordObject):
    """
    Document-level parameters and settings.
    """

    def __init__(self, attributes: dict) -> None:
        super().__init__()
        self.attributes = attributes
        self.default_style = None
        self.sections = None

    def add_to_document(self, word_document):
        self._prepare_document(word_document=word_document)
        self._prepare_sections()
        self._apply_attributes()
        self._add_page_numbering()

    def _prepare_sections(self):
        self.sections = self.document.sections

    def _add_page_numbering(self):
        if ATTRIBUTE_NODES.PAGE_NUMBERING in self.attributes:
            if not self.attributes[ATTRIBUTE_NODES.PAGE_NUMBERING]:
                pass

            elif self.attributes[ATTRIBUTE_NODES.SKIP_COVER_PAGE]:
                add_page_number_skip_cover(
                    self.sections[0].footer.paragraphs[0].add_run()
                )
                self.sections[0].different_first_page_header_footer = True
                sectPr = self.sections[0]._sectPr  # noqa

                pgNumType = OxmlElement("w:pgNumType")
                pgNumType.set(ns.qn("w:start"), "0")
                sectPr.append(pgNumType)

            else:
                add_page_number(self.sections[0].footer.paragraphs[0].add_run())


class ParagraphTextStyle(WordObject):
    """
    Paragraph and text style configuration.
    """

    def __init__(self, attributes: dict, target_style_name: str) -> None:
        super().__init__()
        self.attributes = attributes
        self.target_style_name = target_style_name
        self.target_style = None

    def add_to_document(self, word_document):
        self._prepare_document(word_document=word_document)
        self._prepare_target_style()
        self._apply_text_style_attributes()

    def _prepare_target_style(self):
        self.target_style = self.document.styles[self.target_style_name]

    def _apply_text_style_attributes(self):
        for kwarg, value in self.attributes.items():
            if kwarg in document_text_font_attribute_to_class_map:
                attribute_class = document_text_font_attribute_to_class_map[kwarg]
            elif kwarg in document_text_paragraph_format_attribute_to_class_map:
                attribute_class = document_text_paragraph_format_attribute_to_class_map[
                    kwarg
                ]
            else:
                raise KeyError(f"The key {kwarg} is not a OpenReport key. Skipping.")

            attribute_target = self.target_style
            attribute_class.apply(attribute_target, value)


class ListOfItems(WordObject):
    """
    Base class for document list elements (e.g., table of contents).
    """

    def __init__(self) -> None:
        super().__init__()

    def _prepare_run(self):
        self.this_object = self.document.add_paragraph()
        self.run = self.this_object.add_run()

    def add_to_document(self, word_document):
        pass


class TableOfContents(ListOfItems):
    """
    TableOfContents class
    """

    def __init__(self) -> None:
        super().__init__()

    def add_to_document(self, word_document):
        self._prepare_document(word_document=word_document)
        self._prepare_run()

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = 'TOC \\o "1-5" \\h \\z \\u'

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")

        fldChar3 = OxmlElement("w:t")
        fldChar3.text = "Right-click to update field."

        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")

        r_element = self.run._r  # noqa
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)


class PageBreak(WordObject):
    """
    PageBreak class
    """

    def __init__(self, attributes: dict) -> None:
        super().__init__()
        self.attributes = attributes

    def add_to_document(self, word_document):
        try:
            self._prepare_document(word_document=word_document)

            if not hasattr(self, 'attributes') or not self.attributes:
                raise Exception("OpenReport Error: PageBreak object missing attributes. Please provide number of pages.")

            if ATTRIBUTE_NODES.NUMBER_OF_PAGES not in self.attributes:
                raise Exception("OpenReport Error: PageBreak missing 'number_of_pages' attribute. Please specify how many page breaks to add.")

            num_pages = self.attributes[ATTRIBUTE_NODES.NUMBER_OF_PAGES]

            if not isinstance(num_pages, int):
                raise Exception(f"OpenReport Error: Number of pages must be an integer, got {type(num_pages).__name__}.")

            if num_pages < 0:
                raise Exception(f"OpenReport Error: Number of pages cannot be negative: {num_pages}.")

            if num_pages > 100:
                raise Exception(f"OpenReport Error: Number of pages seems too large: {num_pages}. Please check if this is correct.")

            try:
                for i in range(num_pages + 1):
                    self.document.add_page_break()
            except Exception as e:
                raise Exception(f"OpenReport Error: Failed to add page breaks to document. {str(e)}") from e

        except Exception as e:
            if hasattr(e, 'args') and e.args and "OpenReport Error:" in str(e.args[0]):
                raise
            raise Exception(f"OpenReport Error: Failed to add page break to document. {str(e)}") from e
